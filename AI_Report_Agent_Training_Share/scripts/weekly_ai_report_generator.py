import argparse
import csv
import json
import re
import subprocess
import sys
from datetime import date, datetime, timedelta, timezone
from pathlib import Path

from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.opc.constants import RELATIONSHIP_TYPE as RT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from openpyxl import load_workbook


DEFAULT_INPUT = "weekly_ai_report_input_2026-06-04.json"
DEFAULT_TRACKER = "emerging_ai_tracker.json"
DEFAULT_SUBSCRIBERS = "mailing_list_subscribers.json"
DEFAULT_FORMS_CONFIG = "microsoft_forms_subscription_config.json"
DEFAULT_FORMS_SYNC_STATE = "microsoft_forms_sync_state.json"
DEFAULT_BEST_FOR_MEMORY = "best_for_memory.json"
DEFAULT_SUBSCRIBE_URL = (
    "https://dlrequest/GroupID/Groups/Properties?"
    "identity=ZGUyMjg2Y2MyMjlmNDdiNmIwNTViMzM3YzBhOTU0ZDJ8Z3JvdXA="
)
STATIC_EMAIL_RECIPIENTS = ()
EMAIL_RE = re.compile(r"^[^@\s]+@[^@\s]+\.[^@\s]+$")


def load_input(path: Path) -> dict:
    return json.loads(path.read_text(encoding="utf-8"))


def _normalized_model_key(name: str, maker: str, section: str) -> tuple[str, str, str]:
    return (name.strip().casefold(), maker.strip().casefold(), section.strip().casefold())


def validate_model_research_coverage(data: dict, input_path: Path) -> dict | None:
    """Fail closed unless the plan and completed inquiry ledger cover the live CSV."""
    watchlist = data.get("model_watchlist")
    if not watchlist:
        return None

    csv_name = str(watchlist.get("source_csv") or "").strip()
    plan_name = str(watchlist.get("research_plan") or "").strip()
    ledger_name = str(watchlist.get("research_ledger") or "").strip()
    if not csv_name or not plan_name or not ledger_name:
        raise ValueError(
            "model_watchlist must define source_csv, research_plan, and research_ledger"
        )

    csv_path = Path(csv_name)
    if not csv_path.is_absolute():
        csv_path = input_path.parent / csv_path
    referenced_plan_path = Path(plan_name)
    if not referenced_plan_path.is_absolute():
        referenced_plan_path = input_path.parent / referenced_plan_path
    ledger_path = Path(ledger_name)
    if not ledger_path.is_absolute():
        ledger_path = input_path.parent / ledger_path
    plan_path = (
        referenced_plan_path.with_suffix(".json")
        if referenced_plan_path.suffix.casefold() == ".md"
        else referenced_plan_path
    )
    if not csv_path.exists():
        raise ValueError(f"Model watchlist CSV not found: {csv_path}")
    if not referenced_plan_path.exists():
        raise ValueError(f"Referenced weekly model research plan not found: {referenced_plan_path}")
    if not plan_path.exists():
        raise ValueError(f"Machine-readable weekly model research plan not found: {plan_path}")
    if not ledger_path.exists():
        raise ValueError(f"Completed weekly model inquiry ledger not found: {ledger_path}")

    csv_models: dict[tuple[str, str, str], dict] = {}
    with csv_path.open("r", encoding="utf-8-sig", newline="") as handle:
        reader = csv.DictReader(handle)
        required = {"Name", "Maker", "Section"}
        missing = required - set(reader.fieldnames or [])
        if missing:
            raise ValueError(f"Model watchlist CSV is missing: {', '.join(sorted(missing))}")
        for row in reader:
            name = (row.get("Name") or "").strip()
            maker = (row.get("Maker") or "").strip()
            section = (row.get("Section") or "").strip() or maker
            if name and maker:
                csv_models[_normalized_model_key(name, maker, section)] = {
                    "name": name,
                    "maker": maker,
                    "section": section,
                }

    plan = json.loads(plan_path.read_text(encoding="utf-8"))
    plan_models: dict[tuple[str, str, str], dict] = {}
    duplicate_keys = set()
    for model in plan.get("models", []):
        name = str(model.get("name") or "").strip()
        maker = str(model.get("maker") or "").strip()
        section = str(model.get("section") or "").strip() or maker
        if not name or not maker:
            continue
        key = _normalized_model_key(name, maker, section)
        if key in plan_models:
            duplicate_keys.add(key)
        plan_models[key] = model

    missing_models = sorted(set(csv_models) - set(plan_models))
    extra_models = sorted(set(plan_models) - set(csv_models))
    missing_queries = sorted(
        key for key in csv_models
        if key in plan_models and len(plan_models[key].get("queries") or []) < 2
    )
    report_range = data.get("date_range", {})
    plan_range = plan.get("date_range", {})
    range_matches = (
        plan_range.get("start") == report_range.get("start")
        and plan_range.get("end") == report_range.get("end")
    )
    declared_count = watchlist.get("model_count")
    plan_count = plan.get("model_count")

    ledger = json.loads(ledger_path.read_text(encoding="utf-8"))
    ledger_records: dict[tuple[str, str, str], dict] = {}
    duplicate_ledger_keys = set()
    for record in ledger.get("records", []):
        name = str(record.get("name") or "").strip()
        maker = str(record.get("maker") or "").strip()
        section = str(record.get("section") or "").strip() or maker
        if not name or not maker:
            continue
        key = _normalized_model_key(name, maker, section)
        if key in ledger_records:
            duplicate_ledger_keys.add(key)
        ledger_records[key] = record

    terminal_outcomes = {
        "checked-no-in-window-candidate-surfaced",
        "candidate-reviewed-include",
        "candidate-reviewed-omit",
    }
    missing_ledger_models = sorted(set(csv_models) - set(ledger_records))
    extra_ledger_models = sorted(set(ledger_records) - set(csv_models))
    incomplete_ledger_models = sorted(
        key
        for key in csv_models
        if key in ledger_records
        and (
            not str(ledger_records[key].get("supplemental_query") or "").strip()
            or str(ledger_records[key].get("outcome") or "").strip()
            not in terminal_outcomes
        )
    )
    ledger_range = ledger.get("report_window", {})
    ledger_range_matches = (
        ledger_range.get("start") == report_range.get("start")
        and ledger_range.get("end") == report_range.get("end")
    )

    problems = []
    if declared_count != len(csv_models):
        problems.append(f"input declares {declared_count} models but CSV contains {len(csv_models)}")
    if plan_count != len(csv_models) or len(plan_models) != len(csv_models):
        problems.append(
            f"plan declares {plan_count} and contains {len(plan_models)} unique models; expected {len(csv_models)}"
        )
    if missing_models:
        problems.append(f"{len(missing_models)} CSV model(s) missing from the plan")
    if extra_models:
        problems.append(f"{len(extra_models)} plan model(s) absent from the CSV")
    if duplicate_keys:
        problems.append(f"{len(duplicate_keys)} duplicate model target(s) in the plan")
    if missing_queries:
        problems.append(f"{len(missing_queries)} model target(s) lack both required search queries")
    if not range_matches:
        problems.append("plan date range does not match the report date range")
    if ledger.get("live_target_count") != len(csv_models):
        problems.append(
            f"ledger declares {ledger.get('live_target_count')} targets; expected {len(csv_models)}"
        )
    if len(ledger_records) != len(csv_models):
        problems.append(
            f"ledger contains {len(ledger_records)} unique targets; expected {len(csv_models)}"
        )
    if missing_ledger_models:
        problems.append(f"{len(missing_ledger_models)} CSV model(s) missing from the ledger")
    if extra_ledger_models:
        problems.append(f"{len(extra_ledger_models)} ledger model(s) absent from the CSV")
    if duplicate_ledger_keys:
        problems.append(f"{len(duplicate_ledger_keys)} duplicate target(s) in the ledger")
    if incomplete_ledger_models:
        problems.append(
            f"{len(incomplete_ledger_models)} ledger target(s) lack a terminal inquiry outcome"
        )
    if ledger.get("search_error_count") not in (0, None):
        problems.append(f"ledger reports {ledger.get('search_error_count')} search error(s)")
    if not ledger_range_matches:
        problems.append("ledger date range does not match the report date range")
    if problems:
        raise ValueError("Incomplete weekly model research coverage: " + "; ".join(problems))

    return {
        "csv_path": csv_path,
        "plan_path": plan_path,
        "referenced_plan_path": referenced_plan_path,
        "ledger_path": ledger_path,
        "model_count": len(csv_models),
        "inquiry_count": len(ledger_records),
    }


def _normalized_subject_key(entry: dict) -> str:
    explicit_key = str(entry.get("subject_key") or "").strip()
    subject = explicit_key or str(entry.get("tool") or entry.get("model") or "").strip()
    return re.sub(r"[^a-z0-9]+", " ", subject.casefold()).strip()


def ordered_report_sections(data: dict) -> list[dict]:
    """Keep NVIDIA company news at the top while preserving all other section order."""
    sections = list(data.get("sections", []))
    return sorted(
        sections,
        key=lambda section: 0 if section.get("title") == "NVIDIA AI News" else 1,
    )


def validate_section_structure(data: dict) -> dict | None:
    """Keep NVIDIA news first, preserve the original lanes, and prevent duplicates."""
    if not data.get("section_policy"):
        return None

    required_titles = [
        "NVIDIA AI News",
        "Watchlist Updates",
        "Technological Advancements in AI",
        "New AI Models and Rumors",
    ]
    sections = data.get("sections", [])
    title_counts = {title: 0 for title in required_titles}
    for section in sections:
        title = section.get("title")
        if title in title_counts:
            title_counts[title] += 1

    missing = [title for title, count in title_counts.items() if count == 0]
    repeated = [title for title, count in title_counts.items() if count > 1]
    if missing or repeated:
        problems = []
        if missing:
            problems.append("missing original section(s): " + ", ".join(missing))
        if repeated:
            problems.append("repeated original section(s): " + ", ".join(repeated))
        raise ValueError("Invalid weekly report section structure: " + "; ".join(problems))

    ordered_titles = [str(section.get("title") or "") for section in sections]
    if ordered_titles[:len(required_titles)] != required_titles:
        raise ValueError(
            "Invalid weekly report section structure: NVIDIA AI News must be first, "
            "followed by Watchlist Updates, Technological Advancements in AI, and "
            "New AI Models and Rumors"
        )

    subjects: dict[str, tuple[str, str]] = {}
    duplicates = []
    section_counts = {}
    for section in sections:
        title = str(section.get("title") or "Updates")
        entries = list(section.get("items", [])) + list(section.get("frontier_models", []))
        section_counts[title] = len(entries)
        for entry in entries:
            key = _normalized_subject_key(entry)
            label = str(entry.get("tool") or entry.get("model") or "").strip()
            if not key:
                continue
            prior = subjects.get(key)
            if prior and prior[0] != title:
                duplicates.append(f"{label} ({prior[0]} and {title})")
            else:
                subjects[key] = (title, label)

    if duplicates:
        raise ValueError(
            "Dedicated report subjects must appear in exactly one section: "
            + "; ".join(sorted(set(duplicates)))
        )

    return {"section_counts": section_counts, "unique_subjects": len(subjects)}


def validate_report_depth(data: dict, section_structure: dict | None) -> dict | None:
    policy = data.get("depth_policy")
    if not policy:
        return None

    minimum_subjects = int(policy.get("minimum_unique_subjects", 0))
    minimum_updates = int(policy.get("minimum_update_bullets_per_item", 0))
    section_minimums = policy.get("minimum_items_by_section", {})
    coverage_exception = str(data.get("coverage_exception") or "").strip()
    problems = []

    unique_subjects = (section_structure or {}).get("unique_subjects", 0)
    if unique_subjects < minimum_subjects and not coverage_exception:
        problems.append(f"{unique_subjects} unique subjects; minimum is {minimum_subjects}")

    actual_counts = (section_structure or {}).get("section_counts", {})
    for title, minimum in section_minimums.items():
        actual = int(actual_counts.get(title, 0))
        if actual < int(minimum) and not coverage_exception:
            problems.append(f"{title} has {actual} items; minimum is {minimum}")

    shallow_items = []
    for section in data.get("sections", []):
        for item in section.get("items", []):
            update_count = len(item.get("updates", []))
            depth_exception = str(item.get("depth_exception") or "").strip()
            if update_count < minimum_updates and not depth_exception:
                shallow_items.append(f"{item.get('tool', '(unnamed)')} has {update_count}")
    if shallow_items:
        problems.append(
            f"items below the {minimum_updates}-update minimum: " + ", ".join(shallow_items)
        )

    if problems:
        raise ValueError("Weekly report depth requirements not met: " + "; ".join(problems))

    return {
        "unique_subjects": unique_subjects,
        "minimum_subjects": minimum_subjects,
        "minimum_updates": minimum_updates,
        "coverage_exception": coverage_exception,
    }


def write_json(path: Path, payload: dict) -> None:
    path.write_text(json.dumps(payload, indent=2, ensure_ascii=True) + "\n", encoding="utf-8")


def normalize_email(email: str) -> str:
    return email.strip().lower()


def load_subscribers(path: Path) -> dict:
    if path.exists():
        return json.loads(path.read_text(encoding="utf-8"))
    return {
        "list_name": "Weekly AI Report Mailing List",
        "subscribers": []
    }


def save_subscribers(path: Path, payload: dict) -> None:
    payload["subscribers"] = sorted(
        payload["subscribers"],
        key=lambda entry: (entry["email"], entry.get("name", ""))
    )
    write_json(path, payload)


def load_forms_config(path: Path) -> dict:
    if path.exists():
        return json.loads(path.read_text(encoding="utf-8"))
    return {
        "form_url": "",
        "responses_workbook_url": "",
        "local_synced_workbook_path": "",
        "unsubscribe_workbook_url": "",
        "local_unsubscribe_workbook_path": "",
        "accepted_extensions": [".csv", ".xlsx"],
        "field_map": {
            "email": ["Email", "Email Address", "Work email", "Work Email"],
            "name": ["Name", "Full Name"],
            "team": ["Team", "Organization", "Org", "Department"],
            "timestamp": ["Completion time", "End time", "Submitted At", "Start time"]
        }
    }


def load_forms_sync_state(path: Path) -> dict:
    if path.exists():
        return json.loads(path.read_text(encoding="utf-8"))
    return {
        "version": 1,
        "sources": {
            "subscribe": {
                "file_path": "",
                "processed_rows": 0,
                "header_signature": "",
                "events_by_email": {}
            },
            "unsubscribe": {
                "file_path": "",
                "processed_rows": 0,
                "header_signature": "",
                "events_by_email": {}
            }
        }
    }


def load_best_for_memory(path: Path) -> dict:
    if path.exists():
        return json.loads(path.read_text(encoding="utf-8"))
    return {"items": {}}


def list_subscribers(path: Path) -> int:
    payload = load_subscribers(path)
    print(f"Mailing list: {payload.get('list_name', 'Weekly AI Report Mailing List')}")
    if not payload["subscribers"]:
        print("No subscribers yet.")
        return 0

    for entry in payload["subscribers"]:
        if entry.get("name"):
            print(f"- {entry['name']} <{entry['email']}>")
        else:
            print(f"- {entry['email']}")
    return 0


def add_subscriber(path: Path, email: str, name: str | None = None) -> int:
    payload = load_subscribers(path)
    normalized = normalize_email(email)
    existing = {entry["email"] for entry in payload["subscribers"]}
    if normalized in existing:
        print(f"Subscriber already exists: {normalized}")
        return 0

    payload["subscribers"].append(
        {
            "email": normalized,
            "name": (name or "").strip()
        }
    )
    save_subscribers(path, payload)
    print(f"Subscribed: {normalized}")
    return 0


def remove_subscriber(path: Path, email: str) -> int:
    payload = load_subscribers(path)
    normalized = normalize_email(email)
    before = len(payload["subscribers"])
    payload["subscribers"] = [
        entry for entry in payload["subscribers"]
        if entry["email"] != normalized
    ]
    save_subscribers(path, payload)
    if len(payload["subscribers"]) == before:
        print(f"Subscriber not found: {normalized}")
    else:
        print(f"Unsubscribed: {normalized}")
    return 0


def import_subscribers(path: Path, import_file: Path) -> int:
    payload = load_subscribers(path)
    existing = {entry["email"] for entry in payload["subscribers"]}
    added = 0
    for raw_line in import_file.read_text(encoding="utf-8").splitlines():
        line = raw_line.strip()
        if not line or line.startswith("#"):
            continue
        if "," in line:
            email_part, name_part = line.split(",", 1)
            email = normalize_email(email_part)
            name = name_part.strip()
        else:
            email = normalize_email(line)
            name = ""
        if not email or email in existing:
            continue
        payload["subscribers"].append({"email": email, "name": name})
        existing.add(email)
        added += 1

    save_subscribers(path, payload)
    print(f"Imported {added} subscriber(s) from {import_file}")
    return 0


def choose_column(row: dict, candidates: list[str]) -> str:
    lowered = {key.strip().lower(): key for key in row.keys()}
    for candidate in candidates:
        match = lowered.get(candidate.strip().lower())
        if match:
            return match
    return ""


def parse_forms_timestamp(value: str) -> datetime | None:
    raw = (value or "").strip()
    if not raw:
        return None

    for parser in (
        lambda item: datetime.fromisoformat(item),
        lambda item: datetime.strptime(item, "%Y-%m-%d %H:%M:%S"),
        lambda item: datetime.strptime(item, "%m/%d/%Y %H:%M:%S"),
    ):
        try:
            parsed = parser(raw)
            if parsed.tzinfo is None:
                return parsed.replace(tzinfo=timezone.utc)
            return parsed.astimezone(timezone.utc)
        except ValueError:
            continue
    return None


def serialize_event(event: dict) -> dict:
    payload = dict(event)
    timestamp = payload.get("timestamp")
    if isinstance(timestamp, datetime):
        payload["timestamp"] = timestamp.isoformat()
    return payload


def deserialize_event(payload: dict) -> dict:
    event = dict(payload)
    timestamp = event.get("timestamp")
    if isinstance(timestamp, str):
        parsed = parse_forms_timestamp(timestamp)
        if parsed is not None:
            event["timestamp"] = parsed
    return event


def read_forms_export(responses_file: Path, start_index: int = 0) -> tuple[list[str], list[dict], int]:
    suffix = responses_file.suffix.lower()
    if suffix == ".csv":
        with responses_file.open("r", encoding="utf-8-sig", newline="") as handle:
            reader = csv.DictReader(handle)
            headers = reader.fieldnames or []
            output = []
            total_rows = 0
            for row in reader:
                if not any(str(value or "").strip() for value in row.values()):
                    continue
                total_rows += 1
                if total_rows <= start_index:
                    continue
                output.append({key: str(value or "") for key, value in row.items()})
            return headers, output, total_rows

    if suffix == ".xlsx":
        workbook = load_workbook(filename=responses_file, read_only=True, data_only=True)
        try:
            worksheet = workbook.active
            iterator = worksheet.iter_rows(values_only=True)
            header_row = next(iterator, None)
            if header_row is None:
                return [], [], 0
            headers = [str(value).strip() if value is not None else "" for value in header_row]
            output = []
            total_rows = 0
            for raw_row in iterator:
                if raw_row is None:
                    continue
                row = {}
                empty = True
                for index, header in enumerate(headers):
                    if not header:
                        continue
                    value = raw_row[index] if index < len(raw_row) else ""
                    if value is None:
                        value = ""
                    else:
                        value = str(value)
                    if value.strip():
                        empty = False
                    row[header] = value
                if empty:
                    continue
                total_rows += 1
                if total_rows <= start_index:
                    continue
                output.append(row)
        finally:
            workbook.close()

        return headers, output, total_rows

    raise ValueError(f"Unsupported response file type: {responses_file.suffix}")


def read_forms_rows(responses_file: Path) -> list[dict]:
    _, rows, _ = read_forms_export(responses_file)
    return rows


def build_header_signature(headers: list[str]) -> str:
    return "|".join(header.strip().lower() for header in headers if header.strip())


def item_key(item: dict) -> str:
    return f"{item.get('tool', '').strip()}|{item.get('company', '').strip()}".lower()


def item_change_fingerprint(item: dict) -> str:
    updates = " ".join(update.get("text", "") for update in item.get("updates", []))
    return " ".join(
        [
            item.get("tool", ""),
            item.get("company", ""),
            item.get("summary", ""),
            updates,
        ]
    ).strip().lower()


def tokenize(text: str) -> set[str]:
    return {token for token in re.findall(r"[a-z0-9]+", text.lower()) if len(token) > 2}


def is_significant_change(current_item: dict, previous_fingerprint: str) -> bool:
    current_tokens = tokenize(item_change_fingerprint(current_item))
    previous_tokens = tokenize(previous_fingerprint)
    if not current_tokens or not previous_tokens:
        return True
    overlap = len(current_tokens & previous_tokens)
    union = len(current_tokens | previous_tokens)
    similarity = overlap / union if union else 0
    return similarity < 0.55


def import_from_microsoft_forms(subscribers_path: Path, responses_file: Path, forms_config_path: Path) -> int:
    config = load_forms_config(forms_config_path)
    suffix = responses_file.suffix.lower()
    if suffix not in {ext.lower() for ext in config.get("accepted_extensions", [".csv"])}:
        print(f"Unsupported file type for Forms import: {responses_file.suffix}", file=sys.stderr)
        return 1

    payload = load_subscribers(subscribers_path)
    existing = {entry["email"] for entry in payload["subscribers"]}
    added = 0
    updated = 0

    rows = read_forms_rows(responses_file)

    if not rows:
        print(f"No responses found in {responses_file}")
        return 0

    first_row = rows[0]
    field_map = config.get("field_map", {})
    email_key = choose_column(first_row, field_map.get("email", []))
    name_key = choose_column(first_row, field_map.get("name", []))
    team_key = choose_column(first_row, field_map.get("team", []))

    if not email_key:
        print(
            "Could not find an email column in the Microsoft Forms export. "
            "Update microsoft_forms_subscription_config.json field_map.email to match your form headers.",
            file=sys.stderr,
        )
        return 1

    for row in rows:
        email = normalize_email(row.get(email_key, ""))
        if not email:
            continue
        if not EMAIL_RE.match(email):
            continue

        matched_entry = next((entry for entry in payload["subscribers"] if entry["email"] == email), None)
        name = row.get(name_key, "").strip() if name_key else ""
        team = row.get(team_key, "").strip() if team_key else ""

        if matched_entry is None:
            payload["subscribers"].append(
                {
                    "email": email,
                    "name": name,
                    "team": team,
                    "source": "microsoft_forms",
                    "submitted_at": datetime.now(timezone.utc).isoformat()
                }
            )
            existing.add(email)
            added += 1
        else:
            if name:
                matched_entry["name"] = name
            if team:
                matched_entry["team"] = team
            matched_entry["updated_at"] = datetime.now(timezone.utc).isoformat()
            matched_entry["source"] = "microsoft_forms"
            updated += 1

    save_subscribers(subscribers_path, payload)
    print(f"Imported from Microsoft Forms: added {added}, updated {updated}")
    print(f"Source file: {responses_file}")
    return 0


def import_unsubscribes_from_microsoft_forms(subscribers_path: Path, responses_file: Path, forms_config_path: Path) -> int:
    config = load_forms_config(forms_config_path)
    suffix = responses_file.suffix.lower()
    if suffix not in {ext.lower() for ext in config.get("accepted_extensions", [".csv"])}:
        print(f"Unsupported file type for Forms import: {responses_file.suffix}", file=sys.stderr)
        return 1

    payload = load_subscribers(subscribers_path)
    rows = read_forms_rows(responses_file)

    if not rows:
        print(f"No responses found in {responses_file}")
        return 0

    first_row = rows[0]
    field_map = config.get("field_map", {})
    email_key = choose_column(first_row, field_map.get("email", []))

    if not email_key:
        print(
            "Could not find an email column in the Microsoft Forms unsubscribe export. "
            "Update microsoft_forms_subscription_config.json field_map.email to match your form headers.",
            file=sys.stderr,
        )
        return 1

    unsubscribe_emails = set()
    for row in rows:
        email = normalize_email(row.get(email_key, ""))
        if email and EMAIL_RE.match(email):
            unsubscribe_emails.add(email)

    before = len(payload["subscribers"])
    payload["subscribers"] = [
        entry for entry in payload["subscribers"]
        if entry.get("email") not in unsubscribe_emails
    ]
    removed = before - len(payload["subscribers"])
    save_subscribers(subscribers_path, payload)
    print(f"Imported unsubscribe requests: removed {removed}")
    print(f"Source file: {responses_file}")
    return 0


def should_reset_source_state(source_state: dict, responses_file: Path, current_header_signature: str, total_rows: int) -> bool:
    if not source_state.get("file_path"):
        return True
    if str(responses_file) != source_state.get("file_path"):
        return True
    if current_header_signature != source_state.get("header_signature", ""):
        return True
    if total_rows < int(source_state.get("processed_rows", 0)):
        return True
    return False


def build_events_from_rows(
    rows: list[dict],
    email_key: str,
    name_key: str,
    team_key: str,
    timestamp_key: str,
    event_type: str,
    responses_file: Path,
) -> dict[str, dict]:
    events: dict[str, dict] = {}
    for row in rows:
        email = normalize_email(row.get(email_key, ""))
        if not email or not EMAIL_RE.match(email):
            continue

        timestamp = parse_forms_timestamp(row.get(timestamp_key, "")) if timestamp_key else None
        if timestamp is None:
            timestamp = datetime.now(timezone.utc)

        existing = events.get(email)
        if existing and existing["timestamp"] >= timestamp:
            continue

        events[email] = {
            "email": email,
            "name": row.get(name_key, "").strip() if name_key else "",
            "team": row.get(team_key, "").strip() if team_key else "",
            "timestamp": timestamp,
            "event_type": event_type,
            "source_file": str(responses_file),
            "source_column": email_key,
        }
    return events


def collect_latest_forms_events(
    responses_file: Path,
    forms_config_path: Path,
    event_type: str,
) -> tuple[dict[str, dict], str]:
    config = load_forms_config(forms_config_path)
    suffix = responses_file.suffix.lower()
    if suffix not in {ext.lower() for ext in config.get("accepted_extensions", [".csv"])}:
        raise ValueError(f"Unsupported file type for Forms import: {responses_file.suffix}")

    rows = read_forms_rows(responses_file)
    if not rows:
        return {}, ""

    first_row = rows[0]
    field_map = config.get("field_map", {})
    email_key = choose_column(first_row, field_map.get("email", []))
    name_key = choose_column(first_row, field_map.get("name", []))
    team_key = choose_column(first_row, field_map.get("team", []))
    timestamp_key = choose_column(first_row, field_map.get("timestamp", []))

    if not email_key:
        raise ValueError(
            "Could not find an email column in the Microsoft Forms export. "
            "Update microsoft_forms_subscription_config.json field_map.email to match your form headers."
        )

    events = build_events_from_rows(
        rows,
        email_key,
        name_key,
        team_key,
        timestamp_key,
        event_type,
        responses_file,
    )
    return events, email_key


def collect_latest_forms_events_incremental(
    responses_file: Path,
    forms_config_path: Path,
    event_type: str,
    source_state: dict,
) -> tuple[dict[str, dict], str, dict, int]:
    config = load_forms_config(forms_config_path)
    suffix = responses_file.suffix.lower()
    if suffix not in {ext.lower() for ext in config.get("accepted_extensions", [".csv"])}:
        raise ValueError(f"Unsupported file type for Forms import: {responses_file.suffix}")

    headers, new_rows, total_rows = read_forms_export(
        responses_file,
        start_index=int(source_state.get("processed_rows", 0)),
    )
    current_header_signature = build_header_signature(headers)
    if should_reset_source_state(source_state, responses_file, current_header_signature, total_rows):
        headers, new_rows, total_rows = read_forms_export(responses_file, start_index=0)
        source_state = {
            "file_path": str(responses_file),
            "processed_rows": 0,
            "header_signature": current_header_signature,
            "events_by_email": {}
        }

    if not headers:
        return {}, "", source_state, 0

    first_row = new_rows[0] if new_rows else {header: "" for header in headers}
    if not first_row:
        first_row = {header: "" for header in headers}

    field_map = config.get("field_map", {})
    email_key = choose_column(first_row, field_map.get("email", []))
    if not email_key:
        lowered_headers = {header.strip().lower(): header for header in headers}
        for candidate in field_map.get("email", []):
            matched = lowered_headers.get(candidate.strip().lower())
            if matched:
                email_key = matched
                break
    name_key = choose_column(first_row, field_map.get("name", [])) if first_row else ""
    team_key = choose_column(first_row, field_map.get("team", [])) if first_row else ""
    timestamp_key = choose_column(first_row, field_map.get("timestamp", [])) if first_row else ""

    if not email_key:
        raise ValueError(
            "Could not find an email column in the Microsoft Forms export. "
            "Update microsoft_forms_subscription_config.json field_map.email to match your form headers."
        )

    existing_events = {
        email: deserialize_event(event)
        for email, event in source_state.get("events_by_email", {}).items()
    }
    incoming_events = build_events_from_rows(
        new_rows,
        email_key,
        name_key,
        team_key,
        timestamp_key,
        event_type,
        responses_file,
    )
    for email, event in incoming_events.items():
        current = existing_events.get(email)
        if current is None or current["timestamp"] < event["timestamp"]:
            existing_events[email] = event

    updated_state = {
        "file_path": str(responses_file),
        "processed_rows": total_rows,
        "header_signature": current_header_signature,
        "events_by_email": {
            email: serialize_event(event)
            for email, event in existing_events.items()
        }
    }
    return existing_events, email_key, updated_state, len(new_rows)


def sync_forms_mailing_list(
    subscribers_path: Path,
    subscribe_file: Path,
    unsubscribe_file: Path,
    forms_config_path: Path,
    forms_sync_state_path: Path,
) -> int:
    sync_state = load_forms_sync_state(forms_sync_state_path)
    sources = sync_state.setdefault("sources", {})
    subscribe_source_state = sources.setdefault(
        "subscribe",
        {"file_path": "", "processed_rows": 0, "header_signature": "", "events_by_email": {}},
    )
    unsubscribe_source_state = sources.setdefault(
        "unsubscribe",
        {"file_path": "", "processed_rows": 0, "header_signature": "", "events_by_email": {}},
    )

    subscribe_events, subscribe_column, updated_subscribe_state, subscribe_new_rows = collect_latest_forms_events_incremental(
        subscribe_file,
        forms_config_path,
        "subscribe",
        subscribe_source_state,
    )
    unsubscribe_events, unsubscribe_column, updated_unsubscribe_state, unsubscribe_new_rows = collect_latest_forms_events_incremental(
        unsubscribe_file,
        forms_config_path,
        "unsubscribe",
        unsubscribe_source_state,
    )
    sync_state["sources"]["subscribe"] = updated_subscribe_state
    sync_state["sources"]["unsubscribe"] = updated_unsubscribe_state
    write_json(forms_sync_state_path, sync_state)

    all_emails = sorted(set(subscribe_events) | set(unsubscribe_events))
    subscribers = []
    subscribed_count = 0
    unsubscribed_count = 0

    for email in all_emails:
        subscribe_event = subscribe_events.get(email)
        unsubscribe_event = unsubscribe_events.get(email)

        latest_event = None
        if subscribe_event and unsubscribe_event:
            latest_event = (
                subscribe_event
                if subscribe_event["timestamp"] >= unsubscribe_event["timestamp"]
                else unsubscribe_event
            )
        else:
            latest_event = subscribe_event or unsubscribe_event

        if latest_event and latest_event["event_type"] == "subscribe":
            subscribers.append(
                {
                    "email": email,
                    "name": latest_event.get("name", ""),
                    "team": latest_event.get("team", ""),
                    "source": "microsoft_forms",
                    "submitted_at": latest_event["timestamp"].isoformat()
                }
            )
            subscribed_count += 1
        elif latest_event:
            unsubscribed_count += 1

    payload = {
        "list_name": "Weekly AI Report Mailing List",
        "subscribers": subscribers
    }
    save_subscribers(subscribers_path, payload)
    print(
        f"Synchronized mailing list from Microsoft Forms: "
        f"{subscribed_count} subscribed, {unsubscribed_count} unsubscribed"
    )
    print(f"Subscribe file: {subscribe_file} (email column: {subscribe_column})")
    print(f"Unsubscribe file: {unsubscribe_file} (email column: {unsubscribe_column})")
    print(f"New subscribe rows processed: {subscribe_new_rows}")
    print(f"New unsubscribe rows processed: {unsubscribe_new_rows}")
    print(f"Sync state file: {forms_sync_state_path}")
    return 0


def apply_best_for_memory(data: dict, memory_path: Path) -> dict:
    memory = load_best_for_memory(memory_path)
    memory_items = memory.get("items", {})
    for section in data.get("sections", []):
        for item in section.get("items", []):
            key = item_key(item)
            previous = memory_items.get(key)
            if not previous:
                continue
            if not is_significant_change(item, previous.get("fingerprint", "")):
                item["best_for"] = previous.get("best_for", item.get("best_for", ""))
                if previous.get("best_for_source_url"):
                    item["best_for_source"] = {"url": previous["best_for_source_url"]}
    return data


def update_best_for_memory(data: dict, memory_path: Path) -> None:
    memory = load_best_for_memory(memory_path)
    items = memory.setdefault("items", {})
    for section in data.get("sections", []):
        for item in section.get("items", []):
            key = item_key(item)
            items[key] = {
                "tool": item.get("tool", ""),
                "company": item.get("company", ""),
                "best_for": item.get("best_for", ""),
                "best_for_source_url": item.get("best_for_source", {}).get("url", ""),
                "fingerprint": item_change_fingerprint(item),
                "last_seen_date": item.get("date", ""),
            }
    write_json(memory_path, memory)


def parse_report_item_date(value: str, default_year: int) -> date | None:
    raw = (value or "").strip()
    if not raw:
        return None

    for fmt in ("%Y-%m-%d", "%B %d, %Y", "%b %d, %Y"):
        try:
            return datetime.strptime(raw, fmt).date()
        except ValueError:
            continue

    for fmt in ("%B %d", "%b %d"):
        try:
            parsed = datetime.strptime(raw, fmt).date()
            return parsed.replace(year=default_year)
        except ValueError:
            continue

    return None


def filter_data_to_report_window(data: dict) -> dict:
    window_start = parse_iso_date(data["date_range"]["start"])
    window_end = parse_iso_date(data["date_range"]["end"])
    if window_end < window_start:
        raise ValueError("date_range.end must be on or after date_range.start")

    removed = {"items": 0, "frontier_models": 0}
    for section in data.get("sections", []):
        filtered_items = []
        for item in section.get("items", []):
            item_date = parse_report_item_date(item.get("date", ""), window_start.year)
            if item_date is not None and window_start <= item_date <= window_end:
                filtered_items.append(item)
            else:
                removed["items"] += 1
        if "items" in section:
            section["items"] = filtered_items

        filtered_frontier_models = []
        for item in section.get("frontier_models", []):
            item_date = parse_report_item_date(
                item.get("date", "") or item.get("last_updated", ""),
                window_start.year,
            )
            if item_date is not None and window_start <= item_date <= window_end:
                filtered_frontier_models.append(item)
            else:
                removed["frontier_models"] += 1
        if "frontier_models" in section:
            section["frontier_models"] = filtered_frontier_models

    return removed


def count_report_entries(data: dict) -> int:
    total = 0
    for section in data.get("sections", []):
        total += len(section.get("items", []))
        total += len(section.get("frontier_models", []))
    return total


def find_previous_input_path(current_input_path: Path) -> Path | None:
    candidates = sorted(current_input_path.parent.glob("weekly_ai_report_input_*.json"))
    candidates = [path for path in candidates if path.resolve() != current_input_path.resolve()]
    if not candidates:
        return None
    return candidates[-1]


def build_last_week_diff_lines(current_data: dict, previous_data: dict | None) -> list[str]:
    if previous_data is None:
        return ["No prior weekly input file was available for comparison."]

    current_items = {}
    previous_items = {}
    for section in current_data.get("sections", []):
        for item in section.get("items", []):
            current_items[item_key(item)] = item
    for section in previous_data.get("sections", []):
        for item in section.get("items", []):
            previous_items[item_key(item)] = item

    current_keys = set(current_items)
    previous_keys = set(previous_items)
    new_keys = sorted(current_keys - previous_keys)
    dropped_keys = sorted(previous_keys - current_keys)
    changed_keys = []
    for key in sorted(current_keys & previous_keys):
        if is_significant_change(current_items[key], item_change_fingerprint(previous_items[key])):
            changed_keys.append(key)

    lines = []
    if new_keys:
        headline_item = current_items[new_keys[0]]
        summary_text = headline_item["summary"].rstrip(".").lower()
        lines.append(
            f"New this week: {headline_item['tool']} stood out, with {summary_text}."
        )
    if changed_keys:
        headline_item = current_items[changed_keys[0]]
        summary_text = headline_item["summary"].rstrip(".").lower()
        lines.append(
            f"Biggest update: {headline_item['tool']} advanced this week, with {summary_text}."
        )
    if dropped_keys:
        lines.append(
            "Prior entries without confirmed in-window updates were omitted from this week's report."
        )
    if not lines:
        if not current_keys:
            lines.append("No confirmed in-window updates were included this week.")
        else:
            lines.append("The included model set is broadly similar to last week, with no major lineup changes.")
    return lines


def get_recipient_line(
    path: Path, override_recipients: list[str] | None = None
) -> str:
    # Production delivery uses the maintained Exchange distribution list.
    # `path` remains in the signature for compatibility with the legacy
    # subscriber-management CLI commands, but weekly drafts do not read it.
    recipients = override_recipients or list(STATIC_EMAIL_RECIPIENTS)
    normalized = []
    for recipient in recipients:
        value = str(recipient or "").strip()
        if value and value.casefold() not in {
            existing.casefold() for existing in normalized
        }:
            normalized.append(value)
    return "; ".join(normalized)


def build_email_section_lines(section: dict) -> list[str]:
    title = section.get("title", "Updates")
    lines = [title]
    items = section.get("items", [])
    frontier_models = section.get("frontier_models", [])

    if not items and not frontier_models:
        lines.append("- No confirmed in-window updates.")
        return lines

    if title == "Watchlist Updates":
        for item in items:
            lines.append(f"- {item.get('tool', '')}")
        return lines

    if title == "New AI Models and Rumors":
        released_or_mentioned_models = []
        rumor_summaries = []
        for item in items:
            heading_verb = item.get("heading_verb", "").lower()
            status = item.get("status", "").lower()
            if status == "rumor" or "rumor" in heading_verb:
                rumor_summaries.append(
                    f"- {item.get('tool', '')}: {item.get('summary', '')}"
                )
            else:
                released_or_mentioned_models.append(item.get("tool", ""))

        lines.append("New AI models")
        for name in released_or_mentioned_models:
            lines.append(f"- {name}")

        if rumor_summaries:
            lines.append("")
            lines.append("Rumors")
            lines.extend(rumor_summaries)

        if frontier_models:
            lines.append("")
            lines.append("Next-gen frontier models to watch")
            for item in frontier_models:
                lines.append(
                    f"- {item.get('model', '')} - {item.get('company', '')} - rumored {item.get('rumored_timing', '')}"
                )
        return lines

    for item in items:
        lines.append(f"- {item.get('tool', '')}: {item.get('summary', '')}")
    if frontier_models:
        lines.append("")
        lines.append("Next-gen frontier models to watch")
        for item in frontier_models:
            lines.append(
                f"- {item.get('model', '')} - {item.get('company', '')} - rumored {item.get('rumored_timing', '')}"
            )
    return lines


def build_ai_learning_report_body(
    data: dict,
    subscribe_url: str,
    last_week_diff_lines: list[str],
) -> str:
    lines = [
        "Here is the weekly AI for AI Learning Report generated by OpenAI's Codex.",
        "",
        "This week's updates in brief:",
    ]

    email_sections = ordered_report_sections(data)
    for index, section in enumerate(email_sections):
        lines.append("")
        lines.extend(build_email_section_lines(section))
        if index != len(email_sections) - 1:
            lines.extend(["", "----------------------------------------"])

    lines.extend(
        [
            "",
            "Different from last week",
        ]
    )
    lines.extend(f"- {line}" for line in last_week_diff_lines)

    lines.extend(
        [
            "",
            "Subscribe to the AI for AI Weekly Newsletter by joining the distribution list:",
            subscribe_url,
            "",
            "Questions about this report should be directed to the newsletter owner.",
        ]
    )
    return "\n".join(lines)


def parse_iso_date(value: str) -> date:
    return date.fromisoformat(value)


def ensure_style(document: Document, name: str, font_size: int, bold: bool = False) -> None:
    if name in document.styles:
        style = document.styles[name]
    else:
        style = document.styles.add_style(name, WD_STYLE_TYPE.PARAGRAPH)
    style.font.name = "Calibri"
    style.font.size = Pt(font_size)
    style.font.bold = bold
    style.paragraph_format.space_before = Pt(0)
    style.paragraph_format.space_after = Pt(0)


def configure_document(document: Document) -> None:
    section = document.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

    for style_name in ["Normal", "Heading 1", "Heading 2", "List Bullet", "List Bullet 2"]:
        style = document.styles[style_name]
        style.font.name = "Calibri"
        style.font.size = Pt(11 if style_name == "Normal" else 12)
        style.font.bold = style_name in {"Heading 1", "Heading 2"}
        style.paragraph_format.space_before = Pt(0)
        if style_name == "Normal":
            style.paragraph_format.space_after = Pt(6)
        elif style_name in {"Heading 1", "Heading 2"}:
            style.paragraph_format.space_after = Pt(6)
        else:
            style.paragraph_format.space_after = Pt(3)

    ensure_style(document, "Report Title", 16, bold=True)
    ensure_style(document, "Report Subtitle", 11, bold=False)
    ensure_style(document, "Section Title", 16, bold=True)
    document.styles["Report Title"].paragraph_format.space_after = Pt(3)
    document.styles["Report Subtitle"].paragraph_format.space_after = Pt(6)
    document.styles["Section Title"].paragraph_format.space_before = Pt(12)
    document.styles["Section Title"].paragraph_format.space_after = Pt(6)

    document.styles["Report Title"].font.color.rgb = RGBColor(31, 78, 121)
    document.styles["Section Title"].font.color.rgb = RGBColor(31, 78, 121)
    document.styles["Report Subtitle"].font.color.rgb = RGBColor(89, 89, 89)
    document.styles["Heading 1"].font.color.rgb = RGBColor(31, 78, 121)
    document.styles["Heading 2"].font.color.rgb = RGBColor(31, 78, 121)


def add_page_number(paragraph) -> None:
    run = paragraph.add_run()
    fld_char_begin = OxmlElement("w:fldChar")
    fld_char_begin.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = " PAGE "
    fld_char_end = OxmlElement("w:fldChar")
    fld_char_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_char_begin)
    run._r.append(instr_text)
    run._r.append(fld_char_end)


def add_footer(section) -> None:
    footer = section.footer
    paragraph = footer.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    add_page_number(paragraph)


def add_hyperlink(paragraph, text: str, url: str) -> None:
    part = paragraph.part
    r_id = part.relate_to(url, RT.HYPERLINK, is_external=True)

    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)

    new_run = OxmlElement("w:r")
    r_pr = OxmlElement("w:rPr")

    color = OxmlElement("w:color")
    color.set(qn("w:val"), "0563C1")
    r_pr.append(color)

    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    r_pr.append(underline)

    font_size = OxmlElement("w:sz")
    font_size.set(qn("w:val"), "16")
    r_pr.append(font_size)

    font_size_cs = OxmlElement("w:szCs")
    font_size_cs.set(qn("w:val"), "16")
    r_pr.append(font_size_cs)

    new_run.append(r_pr)
    text_element = OxmlElement("w:t")
    text_element.text = text
    new_run.append(text_element)
    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)


def add_source_link(paragraph, url: str, label: str = "Source") -> None:
    open_paren = paragraph.add_run(" (")
    open_paren.font.size = Pt(8)
    add_hyperlink(paragraph, label, url)
    close_paren = paragraph.add_run(")")
    close_paren.font.size = Pt(8)


def add_title_page(document: Document, title: str, subtitle: str) -> None:
    title_paragraph = document.add_paragraph(style="Report Title")
    title_paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    title_paragraph.add_run(title)

    subtitle_paragraph = document.add_paragraph(style="Report Subtitle")
    subtitle_paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    subtitle_paragraph.add_run(subtitle)


def add_note_paragraph(document: Document, text: str) -> None:
    paragraph = document.add_paragraph(style="Normal")
    run = paragraph.add_run(text)
    run.italic = True
    run.font.color.rgb = RGBColor(80, 80, 80)


def add_glance_section(document: Document, items: list[dict]) -> None:
    document.add_paragraph("This Week at a Glance", style="Heading 1")
    for item in items:
        bullet = document.add_paragraph(style="List Bullet")
        bullet.add_run(
            f"{item['tool']} ({item['company']}, {item['heading_verb']} {item['date']}): {item['summary']}"
        )
        add_source_link(bullet, item["summary_source"]["url"])

        sub_bullet = document.add_paragraph(style="List Bullet 2")
        sub_bullet.add_run(f"Best for: {item['best_for']}")
        add_source_link(sub_bullet, item["best_for_source"]["url"])


def add_detail_sections(document: Document, items: list[dict]) -> None:
    for item in items:
        document.add_paragraph(
            f"{item['tool']} | {item['company']} | {item['heading_verb']} {item['date']}",
            style="Heading 1",
        )

        best_for = document.add_paragraph(style="Normal")
        best_for.add_run(f"Best for: {item['best_for']}")
        add_source_link(best_for, item["best_for_source"]["url"])

        summary = document.add_paragraph(style="Normal")
        summary.add_run(f"Summary: {item['summary']}")
        add_source_link(summary, item["summary_source"]["url"])

        document.add_paragraph("Confirmed updates this week:", style="Normal")
        for update in item["updates"]:
            paragraph = document.add_paragraph(style="List Bullet")
            paragraph.add_run(update["text"])
            add_source_link(paragraph, update["source"]["url"])


def add_compact_highlight_section(document: Document, items: list[dict]) -> None:
    for item in items:
        document.add_paragraph(
            f"{item['tool']} | {item['company']} | {item['heading_verb']} {item['date']}",
            style="Heading 1",
        )
        summary = document.add_paragraph(style="Normal")
        summary.add_run(item["summary"])
        add_source_link(summary, item["summary_source"]["url"])

        for update in item["updates"][:1]:
            paragraph = document.add_paragraph(style="List Bullet")
            paragraph.add_run(update["text"])
            add_source_link(paragraph, update["source"]["url"])


def add_frontier_models_section(document: Document, section: dict) -> None:
    items = section.get("frontier_models", [])
    note = section.get("frontier_models_note", "")
    if not items:
        return

    document.add_paragraph("Next-gen frontier models to watch", style="Heading 1")
    intro = document.add_paragraph(style="Normal")
    intro.add_run(
        "This area tracks rumored next-generation frontier models that have not been publicly released yet. "
        "Models can remain here across multiple weeks when meaningful new reporting or signals emerge."
    )
    if note:
        note_paragraph = document.add_paragraph(style="Normal")
        note_paragraph.add_run(note)

    for item in items:
        bullet = document.add_paragraph(style="List Bullet")
        bullet.add_run(
            f"{item['model']} - {item['company']} - rumored {item['rumored_timing']}"
        )
        add_source_link(bullet, item["source"]["url"])

        if item.get("note"):
            note_bullet = document.add_paragraph(style="List Bullet 2")
            note_bullet.add_run(item["note"])
            add_source_link(note_bullet, item["source"]["url"])


def ensure_documents_folder(folder_name: str) -> Path:
    documents_path = Path.home() / "Documents"
    target_folder = documents_path / folder_name
    target_folder.mkdir(parents=True, exist_ok=True)
    return target_folder


def save_document_with_fallback(document: Document, output_path: Path) -> Path:
    try:
        document.save(output_path)
        return output_path
    except PermissionError:
        timestamp = datetime.now().strftime("%Y%m%d-%H%M%S")
        fallback_path = output_path.with_name(
            f"{output_path.stem} - corrected {timestamp}{output_path.suffix}"
        )
        document.save(fallback_path)
        return fallback_path


def create_outlook_draft(subject: str, body: str, attachment_path: Path, recipient_line: str = "") -> bool:
    powershell_script = rf"""
$outlook = New-Object -ComObject Outlook.Application
$mail = $outlook.CreateItem(0)
$mail.Subject = "{subject.replace('"', "'")}"
$mail.Body = "{body.replace('"', "'")}"
$mail.To = "{recipient_line.replace('"', "'")}"
$null = $mail.Attachments.Add("{str(attachment_path)}")
$mail.Save()
"""
    result = subprocess.run(
        ["powershell", "-NoProfile", "-Command", powershell_script],
        capture_output=True,
        text=True,
        check=False,
    )
    return result.returncode == 0


def write_fallback_draft(target_folder: Path, subject: str, body: str, attachment_path: Path, recipient_line: str) -> Path:
    fallback_path = target_folder / "weekly_ai_report_email_draft.txt"
    fallback_path.write_text(
        f"Subject: {subject}\n"
        f"To: {recipient_line}\n\n"
        f"{body}\n\nAttachment:\n{attachment_path}\n",
        encoding="utf-8",
    )
    return fallback_path


def load_tracker(path: Path) -> dict:
    if path.exists():
        return json.loads(path.read_text(encoding="utf-8"))
    return {
        "threshold_weeks": 3,
        "tracked_models": {},
        "up_and_coming_watchlist": []
    }


def sync_tracker(data: dict, tracker_path: Path) -> dict:
    tracker = load_tracker(tracker_path)
    threshold = tracker.get("threshold_weeks", 3)
    current_week_start = data["date_range"]["start"]
    previous_week_start = (parse_iso_date(current_week_start) - timedelta(days=7)).isoformat()

    rumor_items = []
    for section in data["sections"]:
        if section.get("tracker_enabled"):
            rumor_items.extend(section["items"])

    active_names = []
    for item in rumor_items:
        key = item["tool"]
        active_names.append(key)
        tracked = tracker["tracked_models"].get(key, {"consecutive_weeks": 0, "last_seen_week": None})
        if tracked["last_seen_week"] == current_week_start:
            consecutive_weeks = tracked["consecutive_weeks"]
        elif tracked["last_seen_week"] == previous_week_start:
            consecutive_weeks = tracked["consecutive_weeks"] + 1
        else:
            consecutive_weeks = 1

        tracker["tracked_models"][key] = {
            "company": item["company"],
            "consecutive_weeks": consecutive_weeks,
            "last_seen_week": current_week_start,
            "latest_note": item["summary"],
            "latest_source_url": item["summary_source"]["url"]
        }

    for key, tracked in tracker["tracked_models"].items():
        if tracked["consecutive_weeks"] >= threshold and key not in tracker["up_and_coming_watchlist"]:
            tracker["up_and_coming_watchlist"].append(key)

    tracker["up_and_coming_watchlist"] = [
        key for key in tracker["up_and_coming_watchlist"]
        if tracker["tracked_models"].get(key, {}).get("consecutive_weeks", 0) >= threshold
    ]
    tracker["last_processed_week"] = current_week_start
    tracker["active_this_week"] = active_names

    write_json(tracker_path, tracker)
    return tracker


def add_tracker_note(document: Document, tracker: dict, section: dict) -> None:
    note = section.get("tracker_note")
    if note:
        add_note_paragraph(document, note)

    active_this_week = set(tracker.get("active_this_week", []))
    watchlist = [
        item for item in tracker.get("up_and_coming_watchlist", [])
        if item in active_this_week
    ]
    if watchlist:
        paragraph = document.add_paragraph(style="Normal")
        paragraph.add_run("Up-and-coming watchlist: ")
        paragraph.add_run(", ".join(watchlist))
    else:
        add_note_paragraph(
            document,
            "Up-and-coming watchlist: none yet. Models move onto this watchlist after appearing in the rumors area for more than three consecutive weekly reports."
        )


def build_report(data: dict, output_dir: Path, tracker: dict) -> Path:
    document = Document()
    configure_document(document)
    add_footer(document.sections[0])

    add_title_page(
        document,
        data["report_title"],
        f"Date range: {data['date_range']['display']}",
    )

    report_sections = ordered_report_sections(data)
    for index, section in enumerate(report_sections):
        document.add_paragraph(section["title"], style="Section Title")
        if section.get("tracker_enabled"):
            add_tracker_note(document, tracker, section)
        items = section.get("items", [])
        frontier_models = section.get("frontier_models", [])
        if not items and not frontier_models:
            add_note_paragraph(
                document,
                "No confirmed in-window updates were found for this section."
            )
            if index != len(report_sections) - 1:
                document.add_page_break()
            continue
        if section.get("title") == "Technological Advancements in AI":
            add_note_paragraph(
                document,
                "Cross-cutting capability, research, safety, infrastructure, deployment, evaluation, and adoption advances that are not dedicated model-release entries elsewhere in the report."
            )
            if items:
                add_compact_highlight_section(document, items)
                document.add_paragraph()
                add_detail_sections(document, items)
        else:
            if items:
                add_glance_section(document, items)
                document.add_paragraph()
                add_detail_sections(document, items)
        add_frontier_models_section(document, section)
        if index != len(report_sections) - 1:
            document.add_page_break()

    output_path = output_dir / data["output"]["file_name"]
    return save_document_with_fallback(document, output_path)


def parse_args() -> argparse.Namespace:
    parser = argparse.ArgumentParser()
    parser.add_argument("--input", default=DEFAULT_INPUT)
    parser.add_argument("--tracker", default=DEFAULT_TRACKER)
    parser.add_argument("--subscribers", default=DEFAULT_SUBSCRIBERS)
    parser.add_argument("--forms-config", default=DEFAULT_FORMS_CONFIG)
    parser.add_argument("--forms-sync-state", default=DEFAULT_FORMS_SYNC_STATE)
    parser.add_argument("--best-for-memory", default=DEFAULT_BEST_FOR_MEMORY)
    parser.add_argument("--output-dir")
    parser.add_argument("--skip-email-draft", action="store_true")
    parser.add_argument("--draft-style", choices=["standard", "ai_learning_report"], default="standard")
    parser.add_argument("--subscribe-url", default=DEFAULT_SUBSCRIBE_URL)
    parser.add_argument(
        "--email-recipient",
        action="append",
        dest="email_recipients",
        help="Override production To recipients. Repeat for multiple addresses.",
    )
    parser.add_argument("--subscribe")
    parser.add_argument("--subscriber-name")
    parser.add_argument("--unsubscribe")
    parser.add_argument("--list-subscribers", action="store_true")
    parser.add_argument("--import-subscribers")
    parser.add_argument("--import-forms-responses")
    parser.add_argument("--import-unsubscribe-forms-responses")
    parser.add_argument("--sync-forms-mailing-list", action="store_true")
    parser.add_argument("--subscribe-forms-path")
    parser.add_argument("--unsubscribe-forms-path")
    parser.add_argument("--allow-empty-report", action="store_true")
    return parser.parse_args()


def main() -> int:
    args = parse_args()
    input_path = Path(args.input).resolve()
    tracker_path = Path(args.tracker).resolve()
    subscribers_path = Path(args.subscribers).resolve()
    forms_config_path = Path(args.forms_config).resolve()
    forms_sync_state_path = Path(args.forms_sync_state).resolve()
    best_for_memory_path = Path(args.best_for_memory).resolve()

    if args.list_subscribers:
        return list_subscribers(subscribers_path)
    if args.subscribe:
        return add_subscriber(subscribers_path, args.subscribe, args.subscriber_name)
    if args.unsubscribe:
        return remove_subscriber(subscribers_path, args.unsubscribe)
    if args.import_subscribers:
        return import_subscribers(subscribers_path, Path(args.import_subscribers).resolve())
    if args.import_forms_responses:
        return import_from_microsoft_forms(
            subscribers_path,
            Path(args.import_forms_responses).resolve(),
            forms_config_path,
        )
    if args.import_unsubscribe_forms_responses:
        return import_unsubscribes_from_microsoft_forms(
            subscribers_path,
            Path(args.import_unsubscribe_forms_responses).resolve(),
            forms_config_path,
        )
    if args.sync_forms_mailing_list:
        forms_config = load_forms_config(forms_config_path)
        subscribe_forms_path = Path(
            args.subscribe_forms_path or forms_config.get("local_synced_workbook_path", "")
        ).resolve()
        unsubscribe_forms_path = Path(
            args.unsubscribe_forms_path or forms_config.get("local_unsubscribe_workbook_path", "")
        ).resolve()
        return sync_forms_mailing_list(
            subscribers_path,
            subscribe_forms_path,
            unsubscribe_forms_path,
            forms_config_path,
            forms_sync_state_path,
        )

    if not input_path.exists():
        print(f"Input file not found: {input_path}", file=sys.stderr)
        return 1

    data = load_input(input_path)
    try:
        model_coverage = validate_model_research_coverage(data, input_path)
    except (OSError, ValueError, json.JSONDecodeError) as exc:
        print(str(exc), file=sys.stderr)
        return 1
    previous_input_path = find_previous_input_path(input_path)
    previous_data = load_input(previous_input_path) if previous_input_path and previous_input_path.exists() else None
    filter_counts = filter_data_to_report_window(data)
    try:
        section_structure = validate_section_structure(data)
    except ValueError as exc:
        print(str(exc), file=sys.stderr)
        return 1
    try:
        report_depth = validate_report_depth(data, section_structure)
    except ValueError as exc:
        print(str(exc), file=sys.stderr)
        return 1
    if count_report_entries(data) == 0 and not args.allow_empty_report:
        print(
            "No in-window report entries remain after date filtering. "
            "Add researched updates to the input file or rerun with --allow-empty-report.",
            file=sys.stderr,
        )
        return 1
    data = apply_best_for_memory(data, best_for_memory_path)
    tracker = sync_tracker(data, tracker_path)

    if args.output_dir:
        output_dir = Path(args.output_dir).resolve()
        output_dir.mkdir(parents=True, exist_ok=True)
    else:
        output_dir = ensure_documents_folder(data["output"]["folder_name"])

    report_path = build_report(data, output_dir, tracker)
    update_best_for_memory(data, best_for_memory_path)

    draft_subject = data["output"]["draft_subject"]
    if args.draft_style == "ai_learning_report":
        draft_body = build_ai_learning_report_body(
            data,
            args.subscribe_url,
            build_last_week_diff_lines(data, previous_data),
        )
    else:
        draft_body = (
            f"Attached is the weekly AI releases and updates report for "
            f"{data['date_range']['display']}."
        )
    recipient_line = get_recipient_line(
        subscribers_path, args.email_recipients
    )

    draft_created = False
    fallback_path = None
    if not args.skip_email_draft:
        draft_created = create_outlook_draft(draft_subject, draft_body, report_path, recipient_line)
        if not draft_created:
            fallback_path = write_fallback_draft(output_dir, draft_subject, draft_body, report_path, recipient_line)

    print(f"Report created: {report_path}")
    if model_coverage:
        print(
            "Model research coverage verified: "
            f"{model_coverage['model_count']} of {model_coverage['model_count']} CSV targets planned "
            f"and {model_coverage['inquiry_count']} completed inquiries in "
            f"{model_coverage['ledger_path']}"
        )
    if section_structure:
        counts = section_structure["section_counts"]
        print(
            "Section uniqueness verified: "
            f"{section_structure['unique_subjects']} dedicated subjects across "
            + ", ".join(f"{title}={counts.get(title, 0)}" for title in [
                "NVIDIA AI News",
                "Watchlist Updates",
                "Technological Advancements in AI",
                "New AI Models and Rumors",
            ])
        )
    if report_depth:
        print(
            "Report depth verified: "
            f"{report_depth['unique_subjects']} unique subjects "
            f"(minimum {report_depth['minimum_subjects']}), "
            f"at least {report_depth['minimum_updates']} sourced updates per item"
        )
    print(
        "Filtered outside date range: "
        f"{filter_counts['items']} item(s), "
        f"{filter_counts['frontier_models']} frontier model(s)"
    )
    print(f"Tracker updated: {tracker_path}")
    print(
        "Recipient source: "
        + (
            "command-line configuration"
            if args.email_recipients
            else "static production recipients and Exchange distribution list"
        )
    )
    print(f"Email recipients: {recipient_line if recipient_line else '(none)'}")
    if args.skip_email_draft:
        print("Email draft skipped.")
    elif draft_created:
        print("Outlook draft created with attachment.")
    else:
        print("Outlook draft could not be created.")
        print(f"Fallback draft saved: {fallback_path}")

    return 0


if __name__ == "__main__":
    raise SystemExit(main())
